"""
Document template parser and formatter for Codex.

Reads structural metadata (page size, margins, fonts, spacing) from uploaded
resume/cover letter templates (.docx or .pdf), and generates properly formatted
.docx output files that match that exact structure.
"""

import io
import re
from collections import Counter


# ── TEMPLATE READING ──────────────────────────────────────────────────────────

def parse_docx_template(file_bytes: bytes) -> dict:
    """Extract page layout and font metadata from a .docx file."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError:
        return {"error": "python-docx not installed. Run: pip install python-docx"}

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        return {"error": f"Could not read .docx file: {e}"}

    # ── Page dimensions and margins ──
    section = doc.sections[0] if doc.sections else None
    page = {
        "width_inches":    round(section.page_width.inches,    2) if section else 8.5,
        "height_inches":   round(section.page_height.inches,   2) if section else 11.0,
        "margin_top":      round(section.top_margin.inches,    2) if section else 1.0,
        "margin_bottom":   round(section.bottom_margin.inches, 2) if section else 1.0,
        "margin_left":     round(section.left_margin.inches,   2) if section else 1.0,
        "margin_right":    round(section.right_margin.inches,  2) if section else 1.0,
    }

    # ── Font analysis ──
    # Try document Normal style first (most reliable)
    primary_font = "Calibri"
    body_size_pt = 11
    try:
        normal = doc.styles["Normal"]
        if normal.font.name:
            primary_font = normal.font.name
        if normal.font.size:
            body_size_pt = round(normal.font.size.pt)
    except Exception:
        pass

    # Scan runs for font evidence (fallback / confirmation)
    font_counts = Counter()
    size_counts = Counter()
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name:
                font_counts[run.font.name] += 1
            if run.font.size:
                size_counts[round(run.font.size.pt)] += 1

    if font_counts:
        primary_font = font_counts.most_common(1)[0][0]
    if size_counts:
        body_size_pt = size_counts.most_common(1)[0][0]

    # ── Line spacing ──
    line_spacing = 1.0
    for para in doc.paragraphs:
        if para.text.strip() and para.paragraph_format.line_spacing is not None:
            ls = para.paragraph_format.line_spacing
            try:
                ls_val = float(ls)
                # Twips (1 twip = 1/1440 inch; line spacing stored as twips if > 10)
                if ls_val > 10:
                    line_spacing = round(ls_val / 240, 2)
                else:
                    line_spacing = round(ls_val, 2)
            except Exception:
                pass
            if line_spacing > 0:
                break

    # ── Extract text for writing sample ──
    text_lines = [p.text for p in doc.paragraphs if p.text.strip()]
    extracted_text = "\n".join(text_lines)

    return {
        "source_type":  "docx",
        "page":         page,
        "primary_font": primary_font,
        "body_size_pt": body_size_pt,
        "line_spacing": line_spacing if line_spacing > 0 else 1.0,
        "extracted_text": extracted_text[:6000],
        "font_counts":  dict(font_counts.most_common(5)),
    }


def parse_pdf_template(file_bytes: bytes) -> dict:
    """Extract page dimensions and text from a .pdf file."""
    try:
        from pypdf import PdfReader
    except ImportError:
        return {"error": "pypdf not installed. Run: pip install pypdf"}

    try:
        reader = PdfReader(io.BytesIO(file_bytes))
    except Exception as e:
        return {"error": f"Could not read .pdf file: {e}"}

    page0 = reader.pages[0]
    width_pts  = float(page0.mediabox.width)
    height_pts = float(page0.mediabox.height)

    text = ""
    for page in reader.pages:
        t = page.extract_text()
        if t:
            text += t + "\n"

    return {
        "source_type": "pdf",
        "page": {
            "width_inches":  round(width_pts  / 72, 2),
            "height_inches": round(height_pts / 72, 2),
            "margin_top":    1.0,
            "margin_bottom": 1.0,
            "margin_left":   1.0,
            "margin_right":  1.0,
        },
        # PDFs don't expose font metadata easily — use sensible defaults
        "primary_font": "Calibri",
        "body_size_pt": 11,
        "line_spacing": 1.15,
        "extracted_text": text[:6000],
        "font_counts": {},
        "note": "PDF upload: page dimensions detected. Fonts/margins use standard defaults — upload a .docx for exact formatting."
    }


# ── FORMATTED DOCX GENERATION ─────────────────────────────────────────────────

# Known resume section headers (used for detection heuristics)
_RESUME_HEADERS = {
    "experience", "work experience", "professional experience",
    "education", "skills", "technical skills", "core competencies",
    "projects", "summary", "professional summary", "objective",
    "certifications", "achievements", "awards", "publications",
    "contact", "languages", "volunteer", "leadership",
}


def _is_section_header(line: str) -> bool:
    """Return True if this line looks like a resume section header."""
    stripped = line.strip()
    if not stripped or len(stripped) > 60:
        return False
    # All-caps check
    letters = [c for c in stripped if c.isalpha()]
    if letters and all(c.isupper() for c in letters):
        return True
    # Known header names (case-insensitive)
    if stripped.lower().rstrip(":") in _RESUME_HEADERS:
        return True
    return False


def _is_bullet(line: str) -> bool:
    return line.strip().startswith(("- ", "• ", "* ", "· ", "– ", "— "))


def _is_sub_bullet(line: str) -> bool:
    """Return True if line is a sub-bullet (indented bullet for project details)."""
    return bool(re.match(r'^  +[-•*·–—]\s', line))


def _is_project_heading(line: str) -> bool:
    """Return True if line is a project sub-heading within work experience.
    Format: [PROJECT] Project Title — Context"""
    return line.strip().startswith("[PROJECT]")


def _is_job_entry_line(line: str) -> bool:
    """Return True if line looks like 'Company — Title  >>  Date | Location'
    or 'Company — Title' (with em-dash, en-dash, or pipe separator)."""
    if re.search(r'\s[—–]\s', line) and len(line) < 200:
        return True
    return False


def _split_job_entry(line: str):
    """Split a job entry line into (company_role, date_location).
    Expects format: 'Company — Title  >>  Date | Location'
    or 'Company — Title' (no date).
    The >> delimiter separates left from right.
    """
    if ">>" in line:
        parts = line.split(">>", 1)
        return parts[0].strip(), parts[1].strip()
    return line.strip(), ""


def generate_resume_docx(content_text: str, template_meta: dict) -> bytes:
    """
    Generate a formatted resume .docx from AI-generated text.
    Handles full resume structure: contact header, sections, job entries, bullets.
    Company/Role on the left, Date/Location right-aligned on the same line.
    Applies the structural settings from template_meta.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    doc = Document()
    _apply_page_setup(doc, template_meta)

    font_name = template_meta.get("primary_font", "Calibri")
    body_pt   = template_meta.get("body_size_pt", 11)
    line_sp   = template_meta.get("line_spacing", 1.0)

    # Calculate content width for right-aligned tab stop
    page = template_meta.get("page", {})
    content_w_inches = (
        page.get("width_inches", 8.5)
        - page.get("margin_left", 0.5)
        - page.get("margin_right", 0.5)
    )

    # Set Normal style defaults
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(body_pt)

    lines = content_text.strip().split("\n")
    i = 0
    in_skills_section = False
    in_project_context = False  # tracks whether we're inside a [PROJECT] block
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        if not stripped:
            # Skip blank lines — spacing is controlled by paragraph formats
            i += 1
            continue

        # ── Contact / Name header (first 2 lines) ──
        if i <= 1 and not _is_section_header(stripped) and not _is_bullet(stripped):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i == 0:
                # Candidate name — larger, bold
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(stripped)
                run.bold = True
                run.font.name = font_name
                run.font.size = Pt(body_pt + 4)
            else:
                # Contact details line
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(stripped)
                run.font.name = font_name
                run.font.size = Pt(body_pt)

        elif _is_section_header(stripped):
            header_lower = stripped.lower().rstrip(":")
            in_skills_section = header_lower in ("skills", "technical skills", "core competencies")
            in_project_context = False
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(stripped.upper())
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(body_pt)
            _add_bottom_border(p, color="000000")

        elif _is_project_heading(stripped):
            # [PROJECT] Project Title — Context
            proj_title = stripped.replace("[PROJECT]", "").strip()
            in_project_context = True
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Inches(0.25)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(0)
            _set_line_spacing(p, line_sp)
            run = p.add_run(proj_title)
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(body_pt)

        elif _is_sub_bullet(raw) or (in_project_context and _is_bullet(stripped)):
            # Sub-bullet under a project heading — deeper indent
            bullet_text = re.sub(r'^[-•*·–—]\s+', '', stripped)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent       = Inches(0.5)
            p.paragraph_format.first_line_indent  = Inches(-0.2)
            p.paragraph_format.space_after        = Pt(1)
            p.paragraph_format.space_before       = Pt(0)
            _set_line_spacing(p, line_sp)
            run = p.add_run(f"\u2022 {bullet_text}")
            run.font.name = font_name
            run.font.size = Pt(body_pt)

        elif _is_bullet(stripped):
            # Regular bullet (not under a project heading)
            in_project_context = False
            bullet_text = re.sub(r'^[-•*·–—]\s+', '', stripped)

            # Skills section: bold the category name before colon
            if in_skills_section and ":" in bullet_text:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent       = Inches(0.25)
                p.paragraph_format.first_line_indent  = Inches(-0.2)
                p.paragraph_format.space_after        = Pt(1)
                p.paragraph_format.space_before       = Pt(0)
                _set_line_spacing(p, line_sp)
                cat_name, rest = bullet_text.split(":", 1)
                run_bullet = p.add_run("\u2022 ")
                run_bullet.font.name = font_name
                run_bullet.font.size = Pt(body_pt)
                run_cat = p.add_run(cat_name + ":")
                run_cat.bold = True
                run_cat.font.name = font_name
                run_cat.font.size = Pt(body_pt)
                run_skills = p.add_run(rest)
                run_skills.font.name = font_name
                run_skills.font.size = Pt(body_pt)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent       = Inches(0.25)
                p.paragraph_format.first_line_indent  = Inches(-0.2)
                p.paragraph_format.space_after        = Pt(1)
                p.paragraph_format.space_before       = Pt(0)
                _set_line_spacing(p, line_sp)
                run = p.add_run(f"\u2022 {bullet_text}")
                run.font.name = font_name
                run.font.size = Pt(body_pt)

        elif _is_job_entry_line(stripped):
            # Company — Title  >>  Date | Location
            in_project_context = False
            in_skills_section = False
            left_part, right_part = _split_job_entry(stripped)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(0)
            _set_line_spacing(p, line_sp)

            # Add right-aligned tab stop at content width
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(content_w_inches), WD_TAB_ALIGNMENT.RIGHT)

            run_left = p.add_run(left_part)
            run_left.bold = True
            run_left.font.name = font_name
            run_left.font.size = Pt(body_pt)

            if right_part:
                p.add_run("\t")  # Tab character to push to right
                run_right = p.add_run(right_part)
                run_right.bold = True
                run_right.font.name = font_name
                run_right.font.size = Pt(body_pt)

        else:
            # Plain text — also check for skill category format without bullet
            if in_skills_section and ":" in stripped:
                p = doc.add_paragraph()
                p.paragraph_format.space_after  = Pt(1)
                p.paragraph_format.space_before = Pt(0)
                _set_line_spacing(p, line_sp)
                cat_name, rest = stripped.split(":", 1)
                run_cat = p.add_run(cat_name + ":")
                run_cat.bold = True
                run_cat.font.name = font_name
                run_cat.font.size = Pt(body_pt)
                run_skills = p.add_run(rest)
                run_skills.font.name = font_name
                run_skills.font.size = Pt(body_pt)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                _set_line_spacing(p, line_sp)
                run = p.add_run(stripped)
                run.font.name = font_name
                run.font.size = Pt(body_pt)

        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def generate_cover_letter_docx(content_text: str, template_meta: dict) -> bytes:
    """
    Generate a formatted cover letter .docx from AI-generated text.
    Applies the structural settings from template_meta.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    doc = Document()
    _apply_page_setup(doc, template_meta)

    font_name = template_meta.get("primary_font", "Calibri")
    body_pt   = template_meta.get("body_size_pt", 11)
    line_sp   = template_meta.get("line_spacing", 1.15)

    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(body_pt)

    lines = content_text.strip().split("\n")
    # Detect salutation pattern (Dear ...,) and closing (Sincerely, etc.)
    salutation_words = {"dear", "to whom", "hello", "hi"}
    closing_words    = {"sincerely", "regards", "best regards", "thank you",
                        "respectfully", "warm regards", "best", "yours"}

    for raw in lines:
        stripped = raw.strip()

        if not stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            continue

        lower = stripped.lower().rstrip(",:")
        is_salutation = any(lower.startswith(w) for w in salutation_words)
        is_closing    = any(lower.startswith(w) for w in closing_words)

        p = doc.add_paragraph()
        _set_line_spacing(p, line_sp)

        if is_salutation or is_closing:
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(2)
        else:
            p.paragraph_format.space_after = Pt(8)

        run = p.add_run(stripped)
        run.font.name = font_name
        run.font.size = Pt(body_pt)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── HELPERS ───────────────────────────────────────────────────────────────────

def _apply_page_setup(doc, template_meta: dict):
    """Set page dimensions and margins from template metadata."""
    try:
        from docx.shared import Inches
        page    = template_meta.get("page", {})
        section = doc.sections[0]
        section.page_width     = Inches(page.get("width_inches",  8.5))
        section.page_height    = Inches(page.get("height_inches", 11.0))
        section.top_margin     = Inches(page.get("margin_top",    1.0))
        section.bottom_margin  = Inches(page.get("margin_bottom", 1.0))
        section.left_margin    = Inches(page.get("margin_left",   1.0))
        section.right_margin   = Inches(page.get("margin_right",  1.0))
    except Exception:
        pass


def _set_line_spacing(paragraph, spacing: float):
    """Apply line spacing multiple to a paragraph."""
    try:
        from docx.shared import Pt
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = paragraph._p.get_or_add_pPr()
        spacing_el = OxmlElement("w:spacing")
        # Convert multiple to twips (240 twips = single-line spacing)
        line_val = int(round(spacing * 240))
        spacing_el.set(qn("w:line"), str(line_val))
        spacing_el.set(qn("w:lineRule"), "auto")
        existing = pPr.find(qn("w:spacing"))
        if existing is not None:
            existing.attrib.update(spacing_el.attrib)
        else:
            pPr.append(spacing_el)
    except Exception:
        pass


def _add_bottom_border(paragraph, color: str = "000000", size: int = 6):
    """Add a bottom border to a paragraph (section divider effect)."""
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr  = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    str(size))
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), color)
        pBdr.append(bot)
        pPr.append(pBdr)
    except Exception:
        pass


def build_template_prompt_addon(template_meta: dict, doc_type: str) -> str:
    """
    Return a prompt fragment that instructs Claude to format output
    to fit within the structural constraints of the uploaded template.
    """
    if not template_meta:
        return ""

    page = template_meta.get("page", {})
    font = template_meta.get("primary_font", "Calibri")
    size = template_meta.get("body_size_pt", 11)
    w    = page.get("width_inches", 8.5)
    h    = page.get("height_inches", 11.0)
    mt   = page.get("margin_top",    1.0)
    mb   = page.get("margin_bottom", 1.0)
    ml   = page.get("margin_left",   1.0)
    mr   = page.get("margin_right",  1.0)

    content_w = round(w - ml - mr, 2)
    content_h = round(h - mt - mb, 2)

    if doc_type == "resume":
        return f"""
## DOCUMENT FORMAT CONSTRAINTS (from uploaded resume template)
Page: {w}" × {h}" | Margins: {mt}" top, {mb}" bottom, {ml}" left, {mr}" right
Font: {font} {size}pt | Content area: {content_w}" wide × {content_h}" tall

Format ALL output to fit these constraints:
- Use {font} at {size}pt for body text
- Section headers should be ALL CAPS, bold
- Bullet points use standard dash-space format (- text)
- The resume MUST fit on exactly one page at these settings. Do not exceed one page.
- DO NOT include any explanation or commentary — output only the resume content
"""
    else:
        return f"""
## DOCUMENT FORMAT CONSTRAINTS (from uploaded cover letter template)
Page: {w}" × {h}" | Margins: {mt}" top, {mb}" bottom, {ml}" left, {mr}" right
Font: {font} {size}pt | Content area: {content_w}" wide × {content_h}" tall

Format ALL output to fit these constraints:
- Use {font} at {size}pt for body text
- Target: entire letter fits on 1 page at these settings (typically 250–380 words)
- Include: salutation, 3–4 paragraphs, professional closing
- DO NOT include any explanation or commentary — output only the cover letter
"""
